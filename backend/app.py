from flask import Flask, request, jsonify
from flask_cors import CORS
from pymongo import MongoClient
from bson.objectid import ObjectId
from datetime import datetime
import os
import math
from collections import Counter
from dotenv import load_dotenv
import json

# --- New Imports for File Processing & API Calls ---
import requests
import fitz  # PyMuPDF
import docx # python-docx

load_dotenv()

app = Flask(__name__)
CORS(app)

# MongoDB Connection
MONGO_URI = os.getenv('MONGO_URI', 'mongodb://localhost:27017/')
DB_NAME = os.getenv('DB_NAME', 'kmrl_docintel')
GEMINI_API_KEY = os.getenv('GEMINI_API_KEY') # --- IMPORTANT: Add this to your .env file ---

client = MongoClient(MONGO_URI)
db = client[DB_NAME]
documents_collection = db['documents']

# Create indexes
documents_collection.create_index('title')
documents_collection.create_index('department')
documents_collection.create_index('type')
documents_collection.create_index('tags')


# -------------------------
# FILE EXTRACTION FUNCTION
# -------------------------

def extract_text_from_file(file_storage):
    """Extracts raw text from an uploaded file (PDF, DOCX, TXT)."""
    filename = file_storage.filename
    text = ""
    try:
        # Reset file pointer to the beginning
        file_storage.seek(0)
        
        if filename.endswith('.pdf'):
            pdf_document = fitz.open(stream=file_storage.read(), filetype="pdf")
            for page_num in range(len(pdf_document)):
                page = pdf_document.load_page(page_num)
                blocks = page.get_text("blocks")
                blocks.sort(key=lambda b: (b[1], b[0]))
                for b in blocks:
                    text += b[4]
            pdf_document.close()
        elif filename.endswith('.docx'):
            doc = docx.Document(file_storage)
            
            for para in doc.paragraphs:
                text += para.text + "\n"
            
            if doc.tables:
                text += "\n\n--- Extracted Tables ---\n"
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join(cell.text.strip() for cell in row.cells)
                        text += row_text + "\n"
                    text += "------------------------\n"
        elif filename.endswith('.txt'):
            text = file_storage.read().decode('utf-8')
        else:
            return None # Unsupported file type
    except Exception as e:
        print(f"Error extracting text from {filename}: {str(e)}")
        return None
    return text

# -------------------------
# [UPDATED] GEMINI AI ANALYSIS FUNCTION - NOW INCLUDES STATUS
# -------------------------

def analyze_document_with_gemini(text_content):
    """Uses Gemini API to generate summary, tags, and other metadata, including a suggested status."""
    
    if not GEMINI_API_KEY:
        print("GEMINI_API_KEY not found. Returning mock data.")
        return {
            "title": "Mock Title (GEMINI_API_KEY not set)",
            "summary": "This is mock data. Please set your GEMINI_API_KEY in .env to enable AI analysis.",
            "tags": ["error", "mock-data"],
            "department": "Unknown",
            "type": "Unknown",
            "language": "English",
            "status": "review", # Mock status
            "tables_data": []
        }

    # Truncate text to avoid exceeding API limits (e.g., first 15000 chars)
    truncated_text = text_content[:15000]

    api_url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash:generateContent?key={GEMINI_API_KEY}"

    # --- [UPDATED] Define the JSON structure including 'status' ---
    response_schema = {
        "type": "OBJECT",
        "properties": {
            "title": {"type": "STRING"},
            "summary": {"type": "STRING"},
            "tags": {
                "type": "ARRAY",
                "items": {"type": "STRING"}
            },
            "department": {"type": "STRING"},
            "type": {"type": "STRING"},
            # === NEW: STATUS FIELD ===
            "status": {
                "type": "STRING",
                "enum": ["urgent","approved", "review"]
            },
            # =========================
            "tables_data": { 
                "type": "ARRAY",
                "items": {
                    "type": "OBJECT", 
                    "properties": {
                        "caption": {"type": "STRING"},
                        "data": {
                            "type": "ARRAY",
                            "items": {
                                "type": "ARRAY",
                                "items": {"type": "STRING"}
                            }
                        }
                    },
                    "required": ["caption", "data"]
                }
            }
        },
        "required": ["title", "summary", "tags", "department", "type", "status", "tables_data"]
    }

    # --- [UPDATED] Prompt with instructions for status assignment ---
    system_prompt = (
        "You are an expert document analyzer for a large metro rail company (KMRL). "
        "Analyze the provided document text and return ONLY a valid JSON object. "
        "The text is provided with layout preservation, so blocks are in reading order. "
        "Your tasks are: "
        "1. Generate a **detailed summary** of the document, at least 3-4 sentences long. "
        "2. Extract key `tags`. "
        "3. Make an educated guess for the `department` (e.g., 'Operations', 'Engineering', 'Safety'). "
        "4. Make an educated guess for the `type` (e.g., 'Safety Circular', 'Invoice', 'Tender Document'). "
        "5. **Assign an initial document `status` based on its content**: "
        "   - 'urgent': For Incident Reports, Immediate Safety Circulars, or Critical Directives. "
        "   - 'approved': For routine, finalized documents like published Policies, Board Minutes, or generic Training Materials. "
        "   - 'review': If the document is unclear, highly sensitive, or the type is unknown, forcing a human check. "
        "6. If you detect any tables, extract their data. The `tables_data` field MUST be an array of objects, each with a `caption` and `data` (2D array). "
        "   If no tables are found, this field MUST be an empty array `[]`."
    )


    payload = {
        "systemInstruction": {"parts": [{"text": system_prompt}]},
        "contents": [{
            "parts": [{"text": f"Analyze this document: {truncated_text}"}]
        }],
        "generationConfig": {
            "responseMimeType": "application/json",
            "responseSchema": response_schema
        }
    }

    try:
        response = requests.post(api_url, json=payload, headers={'Content-Type': 'application/json'})
        response.raise_for_status() # Raise an error for bad status codes
        
        result = response.json()
        
        generated_json_text = result.get('candidates', [{}])[0].get('content', {}).get('parts', [{}])[0].get('text', '{}')
        
        processed_data = json.loads(generated_json_text)

        print("--- GEMINI ANALYSIS RESULT ---")
        print(json.dumps(processed_data, indent=2))
        print("------------------------------")
        
        processed_data['language'] = "English" # Default
        
        return processed_data
        
    except Exception as e:
        print(f"Error calling Gemini API: {str(e)}")
        # Fallback in case of API error
        return {
            "title": "Error During Analysis",
            "summary": f"An error occurred while analyzing the document: {str(e)}",
            "tags": ["error"],
            "department": "Unknown",
            "type": "Unknown",
            "language": "English",
            "status": "review", # Hardcoded to 'review' only on error
            "tables_data": []
        }

# -------------------------
# SEMANTIC SEARCH (PURE PYTHON)
# -------------------------

def text_to_vector(text):
    """Convert text to word frequency vector"""
    words = text.lower().split()
    return Counter(words)

def cosine_similarity(vec1, vec2):
    """Compute cosine similarity between two vectors"""
    intersection = set(vec1.keys()) & set(vec2.keys())
    numerator = sum(vec1[x] * vec2[x] for x in intersection)

    sum1 = sum(v * v for v in vec1.values())
    sum2 = sum(v * v for v in vec2.values())

    if sum1 == 0 or sum2 == 0:
        return 0.0

    return numerator / math.sqrt(sum1 * sum2)

def semantic_search(query, collection, top_k=10):
    """Lightweight semantic search without ML models"""
    query_vec = text_to_vector(query)

    results = []

    for doc in collection.find({}):
        full_text = (
            (doc.get("title") or "") + " " +
            (doc.get("summary") or "") + " " +
            (doc.get("content") or "") + " " +
            " ".join(doc.get("tags") or [])
        )

        doc_vec = text_to_vector(full_text)
        score = cosine_similarity(query_vec, doc_vec)

        if score > 0:
            doc["similarity"] = score
            doc["_score"] = score
            results.append(doc)

    results.sort(key=lambda x: x["similarity"], reverse=True)

    for r in results:
        r['_id'] = str(r['_id'])

    return results[:top_k]


# -------------------------
#       ROUTES
# -------------------------

@app.route('/api/documents', methods=['GET'])
def get_documents():
    try:
        search_query = request.args.get('search', '')
        department = request.args.get('department', '')
        doc_type = request.args.get('type', '')
        page = int(request.args.get('page', 1))
        limit = int(request.args.get('limit', 10))
        
        query_filter = {}

        if search_query:
            query_filter['$or'] = [
                {'title': {'$regex': search_query, '$options': 'i'}},
                {'summary': {'$regex': search_query, '$options': 'i'}},
                {'tags': {'$regex': search_query, '$options': 'i'}},
                {'content': {'$regex': search_query, '$options': 'i'}}
            ]
        
        if department and department != 'all':
            query_filter['department'] = department
        
        if doc_type and doc_type != 'all-types':
            formatted_type = doc_type.replace('-', ' ').title()
            query_filter['type'] = formatted_type
        
        total = documents_collection.count_documents(query_filter)
        
        skip = (page - 1) * limit
        documents = list(documents_collection.find(query_filter)
                                   .sort('date', -1)
                                   .skip(skip)
                                   .limit(limit))
        
        for doc in documents:
            doc['_id'] = str(doc['_id'])
        
        return jsonify({
            'documents': documents,
            'pagination': {
                'total': total,
                'page': page,
                'limit': limit,
                'pages': (total + limit - 1) // limit
            }
        })
    except Exception as e:
        print(f"Error in get_documents: {str(e)}")
        return jsonify({'error': str(e)}), 500


# --- [UPDATED] UPLOAD ROUTE - Status is now AI-assigned ---
@app.route('/api/documents/upload', methods=['POST'])
def upload_document():
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        # 1. Extract text from the file
        print(f"Processing file: {file.filename}")
        text_content = extract_text_from_file(file)
        
        if text_content is None:
            return jsonify({'error': 'Unsupported file type or error reading file'}), 400
        
        if not text_content.strip():
            return jsonify({'error': 'File appears to be empty'}), 400
            
        print(f"Extracted {len(text_content)} characters.")

        # 2. Analyze text with Gemini (includes status now)
        print("Analyzing document with Gemini AI...")
        processed_data = analyze_document_with_gemini(text_content)
        print("Analysis complete.")

        # 3. Add remaining data (status is already in processed_data)
        processed_data['content'] = text_content # Store the full text
        processed_data['date'] = datetime.now()
        # processed_data['status'] = 'review' # <-- REMOVED: Status is now AI-assigned
        processed_data['source'] = 'uploaded'
        processed_data['starred'] = False
        
        # 4. Insert into database
        result = documents_collection.insert_one(processed_data)
        processed_data['_id'] = str(result.inserted_id)
        
        print(f"Successfully added document {result.inserted_id} to database.")
        return jsonify(processed_data), 201
        
    except Exception as e:
        print(f"Error in upload_document: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/documents/<doc_id>', methods=['DELETE'])
def delete_document(doc_id):
    try:
        result = documents_collection.delete_one({'_id': ObjectId(doc_id)})
        if result.deleted_count == 0:
            return jsonify({'error': 'Document not found'}), 404
        return jsonify({'message': 'Document deleted successfully'}), 200
    except Exception as e:
        print(f"Error in delete_document: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/search/semantic', methods=['POST'])
def search_semantic_route():
    try:
        data = request.get_json()
        query = data.get('query', '')

        if not query:
            return jsonify({'error': 'Query is required'}), 400

        results = semantic_search(query, documents_collection)
        
        return jsonify({'results': results})
    except Exception as e:
        print(f"Error in semantic search: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/documents/<doc_id>', methods=['GET'])
def get_document(doc_id):
    try:
        doc = documents_collection.find_one({'_id': ObjectId(doc_id)})
        if not doc:
            return jsonify({'error': 'Document not found'}), 404
        
        doc['_id'] = str(doc['_id'])
        return jsonify(doc)
    except Exception as e:
        print(f"Error in get_document: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/documents/<doc_id>', methods=['PUT'])
def update_document(doc_id):
    try:
        data = request.get_json()
        update_data = {}

        if 'status' in data:
            update_data['status'] = data['status']
        if 'tags' in data:
            update_data['tags'] = data['tags']
        if 'starred' in data: 
            update_data['starred'] = data['starred'] 
        
        if not update_data: 
            return jsonify({'error': 'No update fields provided'}), 400

        result = documents_collection.update_one(
            {'_id': ObjectId(doc_id)},
            {'$set': update_data}
        )
        
        if result.matched_count == 0:
            return jsonify({'error': 'Document not found'}), 404
        
        return jsonify({'message': 'Document updated successfully'}), 200
    except Exception as e:
        print(f"Error in update_document: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/stats', methods=['GET'])
def get_stats():
    try:
        total_docs = documents_collection.count_documents({})
        urgent_docs = documents_collection.count_documents({'status': 'urgent'})
        today_docs = documents_collection.count_documents({
            'date': {'$gte': datetime.now().replace(hour=0, minute=0, second=0, microsecond=0)}
        })
        
        return jsonify({
            'total_documents': total_docs,
            'urgent_items': urgent_docs,
            'documents_today': today_docs
        })
    except Exception as e:
        print(f"Error in get_stats: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/health', methods=['GET'])
def health_check():
    return jsonify({'status': 'healthy'})


if __name__ == '__main__':
    if not GEMINI_API_KEY:
        print("🚨 Warning: GEMINI_API_KEY environment variable is not set.")
        print("AI features will be disabled, and mock data will be used.")
    app.run(debug=True, port=5000)
